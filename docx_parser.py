#!/usr/bin/env python3
import sys
from docx import Document


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    
    return '\n'.join(full_text)


def parse_docx(file_path):
    doc = Document(file_path)
    content = {
        'paragraphs': [],
        'tables': [],
        'headers': [],
        'footers': []
    }
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            content['paragraphs'].append(paragraph.text)
    
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        content['tables'].append(table_data)
    
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if paragraph.text.strip():
                content['headers'].append(paragraph.text)
        
        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text.strip():
                content['footers'].append(paragraph.text)
    
    return content


def main():
    if len(sys.argv) < 2:
        print('Usage: python docx_parser.py <docx_file_path>')
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    try:
        print('=== Reading DOCX File ===')
        text = read_docx(file_path)
        print(text)
        
        print('\n=== Parsed Content ===')
        content = parse_docx(file_path)
        
        print(f'\nNumber of paragraphs: {len(content["paragraphs"])}')
        print(f'Number of tables: {len(content["tables"])}')
        print(f'Number of headers: {len(content["headers"])}')
        print(f'Number of footers: {len(content["footers"])}')
        
        if content['tables']:
            print('\nTables:')
            for i, table in enumerate(content['tables'], 1):
                print(f'\nTable {i}:')
                for row in table:
                    print(row)
    
    except FileNotFoundError:
        print(f'Error: File not found: {file_path}')
        sys.exit(1)
    except Exception as e:
        print(f'Error: {e}')
        sys.exit(1)


if __name__ == '__main__':
    main()
