#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
解析开题报告docx文档并提取内容
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def extract_docx_content(docx_path):
    """提取docx文档内容"""
    doc = Document(docx_path)
    content = []
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
    return content

def extract_tables(docx_path):
    """提取表格内容"""
    doc = Document(docx_path)
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    return tables_data

if __name__ == "__main__":
    docx_path = "2026届本科毕业设计（论文）开题报告 .docx"

    # 提取文本内容
    content = extract_docx_content(docx_path)
    print("=" * 50)
    print("文档文本内容:")
    print("=" * 50)
    for i, text in enumerate(content[:50], 1):  # 只显示前50段
        print(f"{i}. {text}")

    # 提取表格内容
    tables = extract_tables(docx_path)
    print("\n" + "=" * 50)
    print("文档表格内容:")
    print("=" * 50)
    for i, table in enumerate(tables, 1):
        print(f"\n表格 {i}:")
        for row in table:
            print(row)
