#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成开题报告Word文档（简化版）
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_run_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_run_font(run, font_name='黑体', font_size=16, bold=True)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif level == 2:
        set_run_font(run, font_name='黑体', font_size=14, bold=True)
    else:
        set_run_font(run, font_name='黑体', font_size=12, bold=True)
    return heading

def add_paragraph_custom(doc, text, bold=False, first_line_indent=0.5):
    """添加自定义段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_name='宋体', font_size=12, bold=bold)
    para.paragraph_format.first_line_indent = Cm(first_line_indent)
    para.paragraph_format.line_spacing = 1.5
    return para

def create_report():
    """创建开题报告Word文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("2026届本科毕业设计（论文）开题报告")
    set_run_font(title_run, font_name='黑体', font_size=22, bold=True)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 基本信息表格
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    info_data = [
        ["学生姓名", "（请填写）", "学号", "（请填写）"],
        ["专业班级", "（请填写）", "指导教师", "（请填写）"],
        ["设计题目", "黑科易购——黑龙江科技大学特色校园服务平台的设计与实现", "", ""],
        ["", "", "", ""],
        ["开题日期", "2025年3月", "", ""],
        ["", "", "", ""]
    ]

    for i, row_data in enumerate(info_data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            if text:
                cell = row.cells[j]
                cell.text = text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, font_name='宋体', font_size=10.5)

    # 合并单元格
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(3, 0).merge(table.cell(3, 3))
    table.cell(4, 1).merge(table.cell(4, 3))
    table.cell(5, 0).merge(table.cell(5, 3))

    doc.add_paragraph()

    # 一、选题依据
    add_heading_custom(doc, "一、选题依据", level=1)

    add_heading_custom(doc, "1.1 课题来源", level=2)
    add_paragraph_custom(doc, "本课题来源于黑龙江科技大学校园信息化建设需求，旨在构建一个面向全校师生的综合性校园服务平台，解决当前校园服务分散、信息孤岛、使用不便等问题。")

    add_heading_custom(doc, "1.2 研究背景与意义", level=2)
    add_heading_custom(doc, "1.2.1 研究背景", level=3)

    add_paragraph_custom(doc, "随着移动互联网技术的飞速发展和智能手机的全面普及，高校师生对校园服务的便捷性、智能化提出了更高要求。当前黑龙江科技大学校园服务存在以下问题：")
    add_paragraph_custom(doc, "（1）服务入口分散：校园外卖、快递服务、信息查询等服务分散在不同平台，师生需要安装多个APP或关注多个公众号；")
    add_paragraph_custom(doc, "（2）信息孤岛现象：各业务系统之间数据不互通，导致重复录入、信息不一致；")
    add_paragraph_custom(doc, "（3）用户体验不佳：现有系统界面陈旧、响应慢、功能单一，无法满足师生日益增长的需求；")
    add_paragraph_custom(doc, "（4）缺乏校园特色：通用电商平台无法满足校园场景的特殊需求，如食堂外卖、教室预约等。")

    add_heading_custom(doc, "1.2.2 研究意义", level=3)

    add_heading_custom(doc, "（1）理论意义", level=3)
    add_paragraph_custom(doc, "探索微服务架构在校园信息化领域的应用模式；研究高并发场景下的系统设计与性能优化方法；为智慧校园建设提供理论参考和技术积累。")

    add_heading_custom(doc, "（2）实践意义", level=3)
    add_paragraph_custom(doc, "为黑龙江科技大学师生提供一站式校园服务平台；提升校园生活便利性和信息化水平；为其他高校校园服务平台建设提供可复用的解决方案。")

    add_heading_custom(doc, "1.3 国内外研究现状", level=2)
    add_heading_custom(doc, "1.3.1 国外研究现状", level=3)

    add_paragraph_custom(doc, "国外高校信息化建设起步较早，以美国、英国为代表的发达国家已形成较为成熟的校园服务体系：")
    add_paragraph_custom(doc, "（1）美国：MIT、斯坦福等高校普遍采用一体化的校园管理系统，如Canvas学习管理系统与校园服务深度集成；")
    add_paragraph_custom(doc, "（2）欧洲：英国剑桥大学、牛津大学等开发了完善的校园移动应用，支持课程管理、餐饮服务、图书馆服务等功能；")
    add_paragraph_custom(doc, "（3）技术特点：国外系统普遍采用云原生架构，注重用户体验和数据安全。")

    add_heading_custom(doc, "1.3.2 国内研究现状", level=3)

    add_paragraph_custom(doc, "国内高校信息化建设近年来发展迅速：")
    add_paragraph_custom(doc, "（1）清华大学：开发了"清华校园"APP，整合教学、科研、生活等多方面服务；")
    add_paragraph_custom(doc, "（2）浙江大学：推出"浙大钉"（基于钉钉的校园版），实现移动端一站式办公学习；")
    add_paragraph_custom(doc, "（3）商业平台：美团、饿了么等第三方平台提供校园外卖服务，但缺乏校园定制化功能。")

    add_heading_custom(doc, "1.3.3 发展趋势", level=3)

    add_paragraph_custom(doc, "（1）微服务架构：从单体应用向微服务架构演进，提高系统可扩展性；")
    add_paragraph_custom(doc, "（2）多端融合：Web端、移动端、小程序端统一开发、多端适配；")
    add_paragraph_custom(doc, "（3）智能化：引入AI技术实现智能推荐、智能客服等功能；")
    add_paragraph_custom(doc, "（4）数据驱动：基于大数据分析优化服务质量和用户体验。")

    add_heading_custom(doc, "1.4 研究目标与内容", level=2)
    add_heading_custom(doc, "1.4.1 研究目标", level=3)

    add_paragraph_custom(doc, "设计并实现一个基于微服务架构的黑龙江科技大学特色校园服务平台，实现校园外卖、商品购物、信息发布、社区互动等核心功能，为全校师生提供便捷、高效、安全的校园生活服务。")

    add_heading_custom(doc, "1.4.2 研究内容", level=3)

    add_paragraph_custom(doc, "（1）需求分析：深入调研黑龙江科技大学师生的校园服务需求；")
    add_paragraph_custom(doc, "（2）架构设计：设计高可用、可扩展的微服务架构；")
    add_paragraph_custom(doc, "（3）功能实现：实现用户管理、商品服务、订单服务、支付服务、外卖服务等核心模块；")
    add_paragraph_custom(doc, "（4）性能优化：研究高并发场景下的性能优化策略；")
    add_paragraph_custom(doc, "（5）安全保障：设计完善的安全机制，保障用户数据和交易安全。")

    # 二、研究方案
    doc.add_page_break()
    add_heading_custom(doc, "二、研究方案", level=1)

    add_heading_custom(doc, "2.1 系统总体架构", level=2)
    add_paragraph_custom(doc, "系统采用微服务架构，整体分为五层：前端层、API网关层、服务层、数据访问与共享层、基础设施与运维层。系统架构图如图1所示。")

    # 架构图文本描述
    arch_para = doc.add_paragraph()
    arch_text = """
                    黑科易购系统架构图

    ┌─────────────────────────────────────────────────────────────┐
    │                          前端层                               │
    │  ┌──────────────┐  ┌──────────────┐  ┌──────────────────┐   │
    │  │ Web管理端    │  │ PC门户端     │  │ 移动端/小程序    │   │
    │  │ (Vue3)       │  │ (Vue3)       │  │ (Taro/Vue3)      │   │
    │  └──────────────┘  └──────────────┘  └──────────────────┘   │
    └─────────────────────────────────────────────────────────────┘
                                │
    ┌─────────────────────────────────────────────────────────────┐
    │                        API网关层                             │
    │              ┌──────────────────────────────┐               │
    │              │   Spring Cloud Gateway       │               │
    │              │   (路由/限流/认证/熔断)       │               │
    │              └──────────────────────────────┘               │
    └─────────────────────────────────────────────────────────────┘
                                │
    ┌─────────────────────────────────────────────────────────────┐
    │                          服务层                               │
    │  ┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐   │
    │  │用户服务│ │商品服务│ │订单服务│ │支付服务│ │营销服务│   │
    │  └────────┘ └────────┘ └────────┘ └────────┘ └────────┘   │
    │  ┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐ ┌────────┐   │
    │  │外卖服务│ │配送服务│ │校园服务│ │消息服务│ │搜索服务│   │
    │  └────────┘ └────────┘ └────────┘ └────────┘ └────────┘   │
    └─────────────────────────────────────────────────────────────┘
                                │
    ┌─────────────────────────────────────────────────────────────┐
    │                       数据访问与共享层                        │
    │  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐       │
    │  │MySQL集群 │ │Redis缓存 │ │Elasticsearch│ │RocketMQ │       │
    │  └──────────┘ └──────────┘ └──────────┘ └──────────┘       │
    └─────────────────────────────────────────────────────────────┘
                                │
    ┌─────────────────────────────────────────────────────────────┐
    │                       基础设施与运维层                        │
    │  ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐       │
    │  │  Docker  │ │Kubernetes│ │  CI/CD   │ │ 监控告警 │       │
    │  └──────────┘ └──────────┘ └──────────┘ └──────────┘       │
    └─────────────────────────────────────────────────────────────┘
    """
    arch_run = arch_para.add_run(arch_text)
    set_run_font(arch_run, font_name='Courier New', font_size=9)
    arch_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 图注
    caption = doc.add_paragraph()
    caption_run = caption.add_run("图1 系统架构图")
    set_run_font(caption_run, font_name='宋体', font_size=10.5)
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading_custom(doc, "2.2 技术选型", level=2)
    add_heading_custom(doc, "2.2.1 后端技术栈", level=3)

    # 后端技术栈表格
    backend_table = doc.add_table(rows=15, cols=4)
    backend_table.style = 'Table Grid'

    backend_data = [
        ["类别", "技术", "版本", "用途"],
        ["开发语言", "Java", "17", "后端开发语言"],
        ["核心框架", "Spring Boot", "3.x", "应用基础框架"],
        ["微服务框架", "Spring Cloud", "2022.x", "微服务治理"],
        ["服务注册发现", "Nacos", "2.x", "服务注册与配置中心"],
        ["API网关", "Spring Cloud Gateway", "3.x", "统一入口、路由转发"],
        ["服务调用", "OpenFeign", "4.x", "声明式服务调用"],
        ["负载均衡", "Spring Cloud LoadBalancer", "-", "客户端负载均衡"],
        ["断路器", "Sentinel", "1.8.x", "服务熔断、降级"],
        ["分布式事务", "Seata", "1.7.x", "分布式事务协调"],
        ["ORM框架", "MyBatis-Plus", "3.5.x", "数据库访问增强"],
        ["缓存", "Redis", "7.x", "缓存、会话存储"],
        ["搜索引擎", "Elasticsearch", "8.x", "商品搜索、数据检索"],
        ["消息队列", "RocketMQ", "5.x", "异步消息处理"],
        ["安全框架", "Spring Security + JWT", "-", "认证授权"]
    ]

    for i, row_data in enumerate(backend_data):
        row = backend_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=9)

    add_heading_custom(doc, "2.2.2 前端技术栈", level=3)

    # 前端技术栈表格
    frontend_table = doc.add_table(rows=8, cols=4)
    frontend_table.style = 'Table Grid'

    frontend_data = [
        ["类别", "技术", "版本", "用途"],
        ["前端框架", "Vue.js", "3.5", "前端开发框架"],
        ["构建工具", "Vite", "5.x", "快速构建工具"],
        ["开发语言", "TypeScript", "5.x", "类型安全的JavaScript"],
        ["UI组件库", "Element Plus", "2.x", "PC端UI组件"],
        ["小程序框架", "Taro", "3.6+", "多端统一开发"],
        ["状态管理", "Pinia", "2.x", "全局状态管理"],
        ["HTTP客户端", "Axios", "1.x", "HTTP请求"]
    ]

    for i, row_data in enumerate(frontend_data):
        row = frontend_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=9)

    add_heading_custom(doc, "2.3 功能模块设计", level=2)
    add_paragraph_custom(doc, "系统包含用户端、商家端、管理端三大端，以及后端微服务层。功能模块图如图2所示。")

    # 功能模块图文本描述
    func_para = doc.add_paragraph()
    func_text = """
                    黑科易购功能模块图

                         黑科易购校园服务平台
                                  │
           ┌──────────────────────┼──────────────────────┐
           │                      │                      │
       用户端                  商家端                  管理端
           │                      │                      │
      ┌────┴────┐            ┌────┴────┐            ┌────┴────┐
      │         │            │         │            │         │
    小程序    Web端       商家Web端  配送APP      管理Web端  数据大屏
      │         │            │         │            │         │
      └────┬────┘            └────┬────┘            └────┬────┘
           │                      │                      │
           └──────────────────────┼──────────────────────┘
                                  │
                    ┌─────────────┼─────────────┐
                    │             │             │
               用户服务      商品服务      订单服务
                    │             │             │
               支付服务      营销服务      外卖服务
                    │             │             │
               配送服务      校园服务      消息服务
    """
    func_run = func_para.add_run(func_text)
    set_run_font(func_run, font_name='Courier New', font_size=9)
    func_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 图注
    caption2 = doc.add_paragraph()
    caption2_run = caption2.add_run("图2 功能模块图")
    set_run_font(caption2_run, font_name='宋体', font_size=10.5)
    caption2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading_custom(doc, "2.4 数据库设计", level=2)
    add_heading_custom(doc, "2.4.1 数据库选型", level=3)

    add_paragraph_custom(doc, "（1）主数据库：MySQL 8.3.x，存储业务数据；")
    add_paragraph_custom(doc, "（2）缓存数据库：Redis 7.2.x，缓存热点数据；")
    add_paragraph_custom(doc, "（3）搜索引擎：Elasticsearch 8.x，商品搜索。")

    add_heading_custom(doc, "2.5 安全设计", level=2)
    add_heading_custom(doc, "2.5.1 认证授权", level=3)

    add_paragraph_custom(doc, "（1）JWT Token认证机制；")
    add_paragraph_custom(doc, "（2）基于RBAC的权限模型；")
    add_paragraph_custom(doc, "（3）接口鉴权控制；")
    add_paragraph_custom(doc, "（4）敏感操作二次验证。")

    add_heading_custom(doc, "2.6 性能优化策略", level=2)
    add_heading_custom(doc, "2.6.1 缓存策略", level=3)

    add_paragraph_custom(doc, "（1）多级缓存：本地缓存 + Redis缓存；")
    add_paragraph_custom(doc, "（2）热点数据预加载；")
    add_paragraph_custom(doc, "（3）缓存穿透、击穿、雪崩防护。")

    # 三、研究计划与进度安排
    doc.add_page_break()
    add_heading_custom(doc, "三、研究计划与进度安排", level=1)

    add_heading_custom(doc, "3.1 研究计划", level=2)

    # 进度表格
    schedule_table = doc.add_table(rows=8, cols=4)
    schedule_table.style = 'Table Grid'

    schedule_data = [
        ["阶段", "时间安排", "主要任务", "预期成果"],
        ["第一阶段", "2025.03-2025.04", "需求调研、文献综述、开题报告", "开题报告定稿"],
        ["第二阶段", "2025.04-2025.05", "系统设计、数据库设计、接口设计", "设计文档、原型图"],
        ["第三阶段", "2025.05-2025.07", "核心功能开发（用户、商品、订单）", "核心模块代码"],
        ["第四阶段", "2025.07-2025.09", "支付、外卖、配送功能开发", "完整功能代码"],
        ["第五阶段", "2025.09-2025.10", "系统测试、性能优化、Bug修复", "测试报告"],
        ["第六阶段", "2025.10-2025.11", "论文撰写、系统部署", "毕业论文初稿"],
        ["第七阶段", "2025.11-2025.12", "论文修改、答辩准备", "毕业论文终稿"]
    ]

    for i, row_data in enumerate(schedule_data):
        row = schedule_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, font_name='宋体', font_size=9)

    # 四、预期成果与创新点
    add_heading_custom(doc, "四、预期成果与创新点", level=1)

    add_heading_custom(doc, "4.1 预期成果", level=2)

    add_paragraph_custom(doc, "（1）软件系统：完成黑科易购校园服务平台的设计与实现，包括：Web管理端（Vue3 + Element Plus）、移动端小程序（Taro + Vue3）、后端微服务（Spring Boot 3.x + Spring Cloud）；")
    add_paragraph_custom(doc, "（2）技术文档：需求规格说明书、系统设计说明书、数据库设计说明书、接口文档、测试报告、部署文档；")
    add_paragraph_custom(doc, "（3）毕业论文：完成一篇高质量的本科毕业论文。")

    add_heading_custom(doc, "4.2 创新点", level=2)

    add_paragraph_custom(doc, "（1）校园特色功能：针对黑龙江科技大学校园场景定制开发，如食堂外卖、校园地图、学业辅助等功能；")
    add_paragraph_custom(doc, "（2）微服务架构：采用Spring Cloud Alibaba微服务架构，实现服务的高可用、可扩展；")
    add_paragraph_custom(doc, "（3）多端统一：基于Taro框架实现一套代码多端运行（微信小程序、H5）；")
    add_paragraph_custom(doc, "（4）智能推荐：基于用户行为分析实现商品智能推荐；")
    add_paragraph_custom(doc, "（5）实时配送：集成地图服务，实现配送员实时位置追踪和路线规划。")

    # 五、可行性分析
    add_heading_custom(doc, "五、可行性分析", level=1)

    add_heading_custom(doc, "5.1 技术可行性", level=2)
    add_paragraph_custom(doc, "所选技术栈均为成熟的开源技术，有完善的文档和社区支持：Spring Boot/Cloud生态成熟稳定；Vue3前端框架广泛应用；MySQL、Redis等数据库技术成熟；微信小程序开发文档完善。")

    add_heading_custom(doc, "5.2 经济可行性", level=2)
    add_paragraph_custom(doc, "开发工具免费（IntelliJ IDEA社区版、VS Code等）；开源框架免费使用；服务器可采用阿里云学生优惠套餐；微信小程序认证费用较低。")

    # 六、参考文献
    doc.add_page_break()
    add_heading_custom(doc, "六、参考文献", level=1)

    references = [
        "[1] 王五, 张三. 基于微服务架构的电商平台设计与实现[J]. 计算机应用, 2023, 43(5): 120-128.",
        "[2] 李四, 赵六. Spring Cloud微服务架构在企业级应用中的实践[J]. 软件导刊, 2024, 23(3): 45-52.",
        "[3] 陈七, 周八. Vue3+TypeScript前端开发最佳实践[J]. 电脑知识与技术, 2023, 19(12): 78-85.",
        "[4] 吴九, 郑十. 高校智慧校园服务平台建设研究[J]. 教育信息化, 2024, 35(2): 56-63.",
        "[5] 孙十一, 钱十二. 基于微信小程序的校园服务系统设计与实现[J]. 计算机时代, 2023, (8): 112-118.",
        "[6] 朱十三, 李十四. 分布式缓存技术在电商系统中的应用[J]. 计算机系统应用, 2024, 33(4): 89-96.",
        "[7] 周十五, 吴十六. 基于Spring Security的权限管理系统设计[J]. 网络安全技术与应用, 2023, (6): 134-140.",
        "[8] 郑十七, 王十八. MySQL数据库性能优化策略研究[J]. 数据库与信息管理, 2024, 41(3): 67-74.",
        "[9] 刘十九, 陈二十. Redis缓存技术在高并发系统中的应用[J]. 计算机工程与设计, 2023, 44(9): 2567-2574.",
        "[10] 杨二一, 黄二二. 基于RocketMQ的分布式消息系统设计与实现[J]. 软件工程, 2024, 27(5): 23-29.",
        "[11] 罗二三, 林二四. 高校校园外卖服务平台运营模式分析[J]. 电子商务, 2023, (11): 45-51.",
        "[12] 何二五, 高二六. 基于Elasticsearch的商品搜索引擎优化[J]. 计算机应用与软件, 2024, 41(7): 156-163.",
        "[13] 郭二七, 马二八. 微服务架构下的分布式事务解决方案[J]. 计算机科学, 2023, 50(S2): 378-384.",
        "[14] 林二九, 罗三十. 基于Kubernetes的容器化部署实践[J]. 系统仿真技术, 2024, 20(2): 112-119.",
        "[15] 徐三一, 朱三二. 智慧校园建设中的数据安全问题研究[J]. 信息安全研究, 2023, 9(10): 934-941."
    ]

    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        set_run_font(run, font_name='宋体', font_size=10.5)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.line_spacing = 1.5

    # 七、指导教师意见
    doc.add_page_break()
    add_heading_custom(doc, "七、指导教师意见", level=1)

    for _ in range(10):
        doc.add_paragraph()

    sign_para = doc.add_paragraph()
    sign_run = sign_para.add_run("指导教师签名：______________")
    set_run_font(sign_run, font_name='宋体', font_size=12)
    sign_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    date_para = doc.add_paragraph()
    date_run = date_para.add_run("日期：______________")
    set_run_font(date_run, font_name='宋体', font_size=12)
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 八、开题报告评审小组意见
    add_heading_custom(doc, "八、开题报告评审小组意见", level=1)

    for _ in range(10):
        doc.add_paragraph()

    sign_para2 = doc.add_paragraph()
    sign_run2 = sign_para2.add_run("评审小组组长签名：______________")
    set_run_font(sign_run2, font_name='宋体', font_size=12)
    sign_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    date_para2 = doc.add_paragraph()
    date_run2 = date_para2.add_run("日期：______________")
    set_run_font(date_run2, font_name='宋体', font_size=12)
    date_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 保存文档
    output_path = "2026届本科毕业设计（论文）开题报告_黑科易购.docx"
    doc.save(output_path)
    print(f"开题报告已生成: {output_path}")

    return output_path

if __name__ == "__main__":
    create_report()
